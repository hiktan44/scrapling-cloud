"""Local file parsing (Firecrawl /parse parity): PDF, DOCX, XLSX, HTML, text.

Turns uploaded documents into markdown/text without any network fetch.
"""

from __future__ import annotations

import io

PARSE_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".html", ".htm", ".txt", ".csv", ".md"}


def _parse_pdf(data: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    return "\n\n".join(chunks).strip(), len(reader.pages)


def _parse_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _parse_xlsx(data: bytes, max_rows: int = 2000) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"## {sheet.title}")
        for index, row in enumerate(sheet.iter_rows(values_only=True)):
            if index >= max_rows:
                lines.append(f"... ({sheet.max_row} satırın ilk {max_rows} satırı gösterildi)")
                break
            cells = ["" if cell is None else str(cell) for cell in row]
            if any(cells):
                lines.append(" | ".join(cells))
    workbook.close()
    return "\n".join(lines).strip()


def _parse_html(data: bytes) -> str:
    from markdownify import markdownify as to_markdown

    from .scraper import main_content_html

    html = data.decode("utf-8", errors="replace")
    return to_markdown(main_content_html(html), heading_style="ATX").strip()


def parse_file(filename: str, data: bytes) -> dict:
    """Parse an uploaded document into markdown. Raises ValueError on unsupported type."""
    lowered = (filename or "").lower()
    extension = "." + lowered.rsplit(".", 1)[-1] if "." in lowered else ""
    if extension not in PARSE_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{extension}'. Supported: {', '.join(sorted(PARSE_EXTENSIONS))}")

    pages: int | None = None
    if extension == ".pdf":
        markdown, pages = _parse_pdf(data)
    elif extension == ".docx":
        markdown = _parse_docx(data)
    elif extension in {".xlsx", ".xls"}:
        markdown = _parse_xlsx(data)
    elif extension in {".html", ".htm"}:
        markdown = _parse_html(data)
    else:  # .txt / .csv / .md
        markdown = data.decode("utf-8", errors="replace").strip()

    return {
        "filename": filename,
        "file_type": extension.lstrip("."),
        "pages": pages,
        "characters": len(markdown),
        "markdown": markdown,
    }
